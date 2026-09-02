"""Document loaders for the supported source formats.

Each loader returns plain text plus a short list of ``(heading, level)`` markers
that the chunker uses to stay section-aware. Heavy parsers (PyMuPDF,
python-docx) are imported lazily so the package imports cleanly even when a
given format's dependency is not installed; a clear error is raised only if you
actually try to load that format.
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".html", ".htm", ".md", ".markdown", ".txt"}


@dataclass
class LoadedDocument:
    """Raw text of a source document plus provenance for citations."""

    name: str
    text: str
    version: str  # content hash, used to tie golden cases to a document version
    headings: list[tuple[str, int]] = field(default_factory=list)


def _version_of(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8")).hexdigest()[:12]


def load_document(path: str | Path) -> LoadedDocument:
    """Load one document, dispatching on file extension."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, headings = _load_pdf(path)
    elif suffix == ".docx":
        text, headings = _load_docx(path)
    elif suffix in {".html", ".htm"}:
        text, headings = _load_html(path)
    elif suffix in {".md", ".markdown"}:
        text, headings = _load_markdown(path)
    elif suffix == ".txt":
        text, headings = path.read_text(encoding="utf-8", errors="replace"), []
    else:
        raise ValueError(f"Unsupported document type: {suffix} ({path})")
    return LoadedDocument(name=path.name, text=text, version=_version_of(text), headings=headings)


def load_folder(folder: str | Path) -> list[LoadedDocument]:
    """Load every supported document in a folder (recursively)."""
    folder = Path(folder)
    docs: list[LoadedDocument] = []
    for p in sorted(folder.rglob("*")):
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES:
            docs.append(load_document(p))
    return docs


# --------------------------------------------------------------------------- #
# Format-specific loaders (lazy imports)
# --------------------------------------------------------------------------- #
def _load_pdf(path: Path) -> tuple[str, list[tuple[str, int]]]:
    try:
        import pymupdf as fitz  # modern PyMuPDF module name
    except ImportError:
        try:
            import fitz  # older PyMuPDF exposes the same API as 'fitz'
        except ImportError as exc:  # pragma: no cover
            raise ImportError("PyMuPDF is required to load PDFs: pip install pymupdf") from exc
    parts: list[str] = []
    headings: list[tuple[str, int]] = []
    with fitz.open(path) as doc:
        for page in doc:
            # 'blocks' gives us font sizes we could use to infer headings; here
            # we keep it simple and rely on the chunker's heading heuristics.
            parts.append(page.get_text("text"))
    return "\n".join(parts), headings


def _load_docx(path: Path) -> tuple[str, list[tuple[str, int]]]:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ImportError("python-docx is required to load DOCX: pip install python-docx") from exc
    document = docx.Document(str(path))
    lines: list[str] = []
    headings: list[tuple[str, int]] = []
    for para in document.paragraphs:
        style = (para.style.name or "").lower() if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if style.startswith("heading"):
            level = int(re.sub(r"\D", "", style) or "1")
            headings.append((text, level))
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)
    return "\n".join(lines), headings


def _load_html(path: Path) -> tuple[str, list[tuple[str, int]]]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        # Fall back to a crude tag strip if bs4 is unavailable.
        text = re.sub(r"<[^>]+>", " ", raw)
        return re.sub(r"\s+\n", "\n", text), []
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    headings: list[tuple[str, int]] = []
    for h in soup.find_all(re.compile(r"^h[1-6]$")):
        headings.append((h.get_text(strip=True), int(h.name[1])))
    return soup.get_text("\n"), headings


def _load_markdown(path: Path) -> tuple[str, list[tuple[str, int]]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    headings = [
        (m.group(2).strip(), len(m.group(1)))
        for m in re.finditer(r"^(#{1,6})\s+(.*)$", text, re.M)
    ]
    return text, headings
